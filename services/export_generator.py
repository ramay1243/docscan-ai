from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from io import BytesIO
from datetime import datetime
import logging
import os

logger = logging.getLogger(__name__)

def hex_to_rgb(hex_color):
    """Конвертирует hex цвет в RGB"""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def generate_analysis_word(analysis_data, filename="document.pdf", branding_settings=None):
    """Генерирует Word документ (DOCX) с результатами анализа
    
    Args:
        analysis_data: Данные анализа
        filename: Имя файла
        branding_settings: Настройки брендинга (dict с logo_path, primary_color, secondary_color, company_name)
    """
    try:
        # Получаем настройки брендинга или используем значения по умолчанию
        if branding_settings and branding_settings.get('is_active'):
            primary_color = branding_settings.get('primary_color', '#4361ee')
            secondary_color = branding_settings.get('secondary_color', '#764ba2')
            company_name = branding_settings.get('company_name')
            logo_path = branding_settings.get('logo_path')
        else:
            primary_color = '#4361ee'
            secondary_color = '#764ba2'
            company_name = None
            logo_path = None
        
        doc = Document()
        
        # Настройка шрифта по умолчанию для поддержки кириллицы
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        # Логотип (если есть)
        if logo_path and os.path.exists(logo_path):
            try:
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(logo_path, width=Inches(2))
                doc.add_paragraph()  # Пустая строка после логотипа
            except Exception as e:
                logger.warning(f"⚠️ Не удалось добавить логотип: {e}")
        
        # Заголовок
        title_text = company_name if company_name else 'Анализ документа'
        title = doc.add_heading(title_text, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(24)
        # Конвертируем hex в RGB
        rgb_color = hex_to_rgb(primary_color)
        title_run.font.color.rgb = RGBColor(rgb_color[0], rgb_color[1], rgb_color[2])
        title_run.bold = True
        
        # Информация о документе
        doc.add_paragraph(f'Файл: {filename}')
        doc.add_paragraph(f'Дата анализа: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
        doc.add_paragraph(f'Тип документа: {analysis_data.get("document_type_name", "Не определен")}')
        doc.add_paragraph()  # Пустая строка
        
        # Общий уровень риска
        risk_level = analysis_data.get('executive_summary', {}).get('risk_level', 'LOW')
        risk_colors = {
            'CRITICAL': RGBColor(229, 62, 62),
            'HIGH': RGBColor(221, 107, 32),
            'MEDIUM': RGBColor(214, 158, 46),
            'LOW': RGBColor(56, 161, 105)
        }
        risk_color = risk_colors.get(risk_level, RGBColor(49, 130, 206))
        
        risk_icon = analysis_data.get('executive_summary', {}).get('risk_icon', '⚠️')
        risk_desc = analysis_data.get('executive_summary', {}).get('risk_description', 'Риск не определен')
        
        risk_heading = doc.add_heading(f'{risk_icon} Уровень риска: {risk_level}', 1)
        risk_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        risk_heading_run = risk_heading.runs[0]
        risk_heading_run.font.color.rgb = risk_color
        risk_heading_run.bold = True
        
        doc.add_paragraph(risk_desc).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # Пустая строка
        
        # Решение
        decision = analysis_data.get('executive_summary', {}).get('decision_support', '')
        if decision:
            decision_para = doc.add_paragraph()
            decision_para.add_run('💡 Решение: ').bold = True
            decision_para.add_run(decision)
            decision_para.paragraph_format.space_after = Pt(12)
            doc.add_paragraph()  # Пустая строка
        
        # Статистика рисков
        risk_stats = analysis_data.get('risk_analysis', {}).get('risk_statistics', {})
        if risk_stats:
            doc.add_heading('📊 Статистика рисков', 2)
            
            stats_table = doc.add_table(rows=6, cols=2)
            stats_table.style = 'Light Grid Accent 1'
            
            # Заголовки
            header_cells = stats_table.rows[0].cells
            header_cells[0].text = 'Уровень риска'
            header_cells[1].text = 'Количество'
            for cell in header_cells:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Данные
            stats_data = [
                ('Критических', risk_stats.get('CRITICAL', 0)),
                ('Высоких', risk_stats.get('HIGH', 0)),
                ('Средних', risk_stats.get('MEDIUM', 0)),
                ('Низких', risk_stats.get('LOW', 0)),
                ('Всего', risk_stats.get('total', 0))
            ]
            
            for i, (label, value) in enumerate(stats_data, 1):
                row_cells = stats_table.rows[i].cells
                row_cells[0].text = label
                row_cells[1].text = str(value)
                if i == 5:  # Последняя строка "Всего"
                    for cell in row_cells:
                        cell.paragraphs[0].runs[0].bold = True
            
            doc.add_paragraph()  # Пустая строка
        
        # Экспертная оценка
        expert = analysis_data.get('expert_analysis', {})
        
        legal = expert.get('legal_expertise', '')
        if legal and legal != 'Юридический анализ не выявил критических нарушений':
            doc.add_heading('🧑‍⚖️ Юридическая экспертиза', 2)
            doc.add_paragraph(legal)
            doc.add_paragraph()
        
        financial = expert.get('financial_analysis', '')
        if financial and financial != 'Финансовые условия требуют дополнительной проверки':
            doc.add_heading('💰 Финансовый анализ', 2)
            doc.add_paragraph(financial)
            doc.add_paragraph()
        
        operational = expert.get('operational_risks', '')
        if operational and operational != 'Операционные риски находятся в допустимых пределах':
            doc.add_heading('⚙️ Операционные риски', 2)
            doc.add_paragraph(operational)
            doc.add_paragraph()
        
        strategic = expert.get('strategic_assessment', '')
        if strategic and strategic != 'Документ соответствует базовым стратегическим целям':
            doc.add_heading('🎯 Стратегическая оценка', 2)
            doc.add_paragraph(strategic)
            doc.add_paragraph()
        
        # Ключевые риски
        key_risks = analysis_data.get('risk_analysis', {}).get('key_risks', [])
        if key_risks:
            doc.add_page_break()
            doc.add_heading('⚠️ Детальный анализ рисков', 2)
            
            for i, risk in enumerate(key_risks, 1):
                risk_level = risk.get('level', 'MEDIUM')
                risk_color = risk_colors.get(risk_level, RGBColor(49, 130, 206))
                
                risk_title = doc.add_heading(f"{i}. {risk.get('icon', '⚠️')} {risk.get('title', 'Риск')} ({risk_level})", 3)
                risk_title_run = risk_title.runs[0]
                risk_title_run.font.color.rgb = risk_color
                
                doc.add_paragraph(risk.get('description', ''))
                doc.add_paragraph()
        
        # Рекомендации
        recommendations = analysis_data.get('recommendations', {})
        if recommendations:
            doc.add_page_break()
            doc.add_heading('💡 Практические рекомендации', 2)
            
            practical_actions = recommendations.get('practical_actions', [])
            if practical_actions:
                doc.add_heading('📋 Рекомендуемые действия:', 3)
                for i, action in enumerate(practical_actions, 1):
                    para = doc.add_paragraph(f"{i}. ", style='List Number')
                    if isinstance(action, dict):
                        action_text = action.get('action', action.get('title', ''))
                        effect = action.get('effect', action.get('description', ''))
                        para.add_run(action_text).bold = True
                        if effect:
                            para.add_run(f" - {effect}")
                    else:
                        para.add_run(str(action))
            
            priority_actions = recommendations.get('priority_actions', [])
            if priority_actions:
                doc.add_paragraph()
                doc.add_heading('🚨 Срочные действия:', 3)
                for action in priority_actions:
                    para = doc.add_paragraph("• ", style='List Bullet')
                    if isinstance(action, dict):
                        action_text = action.get('action', action.get('title', ''))
                        para.add_run(action_text)
                    else:
                        para.add_run(str(action))
        
        # Футер
        doc.add_paragraph()
        footer = doc.add_paragraph('Сгенерировано DocScan AI - https://docscan-ai.ru')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.runs[0]
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(108, 117, 125)
        
        # Сохраняем в BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        logger.info(f"✅ Word документ успешно сгенерирован для файла: {filename}")
        return buffer.getvalue()
        
    except Exception as e:
        logger.error(f"❌ Ошибка генерации Word документа: {e}")
        import traceback
        logger.error(traceback.format_exc())
        raise


def generate_analysis_excel(analysis_data, filename="document.pdf", branding_settings=None):
    """Генерирует Excel файл (XLSX) с результатами анализа
    
    Args:
        analysis_data: Данные анализа
        filename: Имя файла
        branding_settings: Настройки брендинга (dict с logo_path, primary_color, secondary_color, company_name)
    """
    try:
        # Получаем настройки брендинга или используем значения по умолчанию
        if branding_settings and branding_settings.get('is_active'):
            primary_color = branding_settings.get('primary_color', '#4361ee')
            secondary_color = branding_settings.get('secondary_color', '#764ba2')
            company_name = branding_settings.get('company_name')
            logo_path = branding_settings.get('logo_path')
        else:
            primary_color = '#4361ee'
            secondary_color = '#764ba2'
            company_name = None
            logo_path = None
        
        wb = Workbook()
        ws = wb.active
        ws.title = company_name if company_name else "Анализ документа"
        
        # Конвертируем hex в RGB для Excel (убираем #)
        primary_color_excel = primary_color.lstrip('#')
        
        # Стили
        header_fill = PatternFill(start_color=primary_color_excel, end_color=primary_color_excel, fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=12)
        title_font = Font(bold=True, size=16, color=primary_color_excel)
        risk_font = Font(bold=True, size=11)
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        center_align = Alignment(horizontal='center', vertical='center')
        wrap_align = Alignment(wrap_text=True, vertical='top')
        
        row = 1
        
        # Заголовок
        title_text = company_name if company_name else "Анализ документа"
        ws.merge_cells(f'A{row}:D{row}')
        cell = ws[f'A{row}']
        cell.value = title_text
        cell.font = title_font
        cell.fill = header_fill
        cell.font = Font(bold=True, color="FFFFFF", size=16)
        cell.alignment = center_align
        row += 1
        
        # Информация о документе
        row += 1
        ws[f'A{row}'] = "Файл:"
        ws[f'B{row}'] = filename
        ws[f'B{row}'].font = Font(bold=True)
        row += 1
        
        ws[f'A{row}'] = "Дата анализа:"
        ws[f'B{row}'] = datetime.now().strftime("%d.%m.%Y %H:%M")
        row += 1
        
        ws[f'A{row}'] = "Тип документа:"
        ws[f'B{row}'] = analysis_data.get('document_type_name', 'Не определен')
        row += 2
        
        # Общий уровень риска
        risk_level = analysis_data.get('executive_summary', {}).get('risk_level', 'LOW')
        risk_colors_map = {
            'CRITICAL': 'E53E3E',
            'HIGH': 'DD6B20',
            'MEDIUM': 'D69E2E',
            'LOW': '38A169'
        }
        risk_color = risk_colors_map.get(risk_level, '3182CE')
        risk_fill = PatternFill(start_color=risk_color, end_color=risk_color, fill_type="solid")
        
        ws.merge_cells(f'A{row}:D{row}')
        cell = ws[f'A{row}']
        risk_icon = analysis_data.get('executive_summary', {}).get('risk_icon', '⚠️')
        risk_desc = analysis_data.get('executive_summary', {}).get('risk_description', 'Риск не определен')
        cell.value = f"{risk_icon} Уровень риска: {risk_level}"
        cell.font = risk_font
        cell.fill = risk_fill
        cell.font = Font(bold=True, color="FFFFFF", size=14)
        cell.alignment = center_align
        row += 1
        
        ws.merge_cells(f'A{row}:D{row}')
        ws[f'A{row}'] = risk_desc
        ws[f'A{row}'].alignment = center_align
        row += 2
        
        # Решение
        decision = analysis_data.get('executive_summary', {}).get('decision_support', '')
        if decision:
            ws[f'A{row}'] = "💡 Решение:"
            ws[f'A{row}'].font = Font(bold=True)
            ws.merge_cells(f'B{row}:D{row}')
            ws[f'B{row}'] = decision
            ws[f'B{row}'].alignment = wrap_align
            row += 2
        
        # Статистика рисков
        risk_stats = analysis_data.get('risk_analysis', {}).get('risk_statistics', {})
        if risk_stats:
            ws[f'A{row}'] = "📊 Статистика рисков"
            ws[f'A{row}'].font = Font(bold=True, size=12)
            row += 1
            
            # Заголовки таблицы
            headers = ['Уровень риска', 'Количество']
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=row, column=col)
                cell.value = header
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = center_align
                cell.border = border
            row += 1
            
            # Данные
            stats_data = [
                ('Критических', risk_stats.get('CRITICAL', 0)),
                ('Высоких', risk_stats.get('HIGH', 0)),
                ('Средних', risk_stats.get('MEDIUM', 0)),
                ('Низких', risk_stats.get('LOW', 0)),
                ('Всего', risk_stats.get('total', 0))
            ]
            
            for label, value in stats_data:
                ws.cell(row=row, column=1).value = label
                ws.cell(row=row, column=2).value = value
                if label == 'Всего':
                    for col in [1, 2]:
                        ws.cell(row=row, column=col).font = Font(bold=True)
                for col in [1, 2]:
                    ws.cell(row=row, column=col).border = border
                    ws.cell(row=row, column=col).alignment = center_align
                row += 1
            
            row += 1
        
        # Экспертная оценка
        expert = analysis_data.get('expert_analysis', {})
        
        sections = [
            ('🧑‍⚖️ Юридическая экспертиза', expert.get('legal_expertise', '')),
            ('💰 Финансовый анализ', expert.get('financial_analysis', '')),
            ('⚙️ Операционные риски', expert.get('operational_risks', '')),
            ('🎯 Стратегическая оценка', expert.get('strategic_assessment', ''))
        ]
        
        skip_texts = [
            'Юридический анализ не выявил критических нарушений',
            'Финансовые условия требуют дополнительной проверки',
            'Операционные риски находятся в допустимых пределах',
            'Документ соответствует базовым стратегическим целям'
        ]
        
        for section_title, section_text in sections:
            if section_text and section_text not in skip_texts:
                ws[f'A{row}'] = section_title
                ws[f'A{row}'].font = Font(bold=True, size=12)
                row += 1
                
                ws.merge_cells(f'A{row}:D{row}')
                ws[f'A{row}'] = section_text
                ws[f'A{row}'].alignment = wrap_align
                row += 2
        
        # Ключевые риски
        key_risks = analysis_data.get('risk_analysis', {}).get('key_risks', [])
        if key_risks:
            ws[f'A{row}'] = "⚠️ Детальный анализ рисков"
            ws[f'A{row}'].font = Font(bold=True, size=12)
            row += 1
            
            for i, risk in enumerate(key_risks, 1):
                risk_level = risk.get('level', 'MEDIUM')
                risk_title = f"{i}. {risk.get('icon', '⚠️')} {risk.get('title', 'Риск')} ({risk_level})"
                
                ws[f'A{row}'] = risk_title
                ws[f'A{row}'].font = Font(bold=True)
                row += 1
                
                ws.merge_cells(f'A{row}:D{row}')
                ws[f'A{row}'] = risk.get('description', '')
                ws[f'A{row}'].alignment = wrap_align
                row += 1
        
        # Рекомендации
        recommendations = analysis_data.get('recommendations', {})
        if recommendations:
            ws[f'A{row}'] = "💡 Практические рекомендации"
            ws[f'A{row}'].font = Font(bold=True, size=12)
            row += 1
            
            practical_actions = recommendations.get('practical_actions', [])
            if practical_actions:
                ws[f'A{row}'] = "📋 Рекомендуемые действия:"
                ws[f'A{row}'].font = Font(bold=True)
                row += 1
                
                for i, action in enumerate(practical_actions, 1):
                    if isinstance(action, dict):
                        action_text = action.get('action', action.get('title', ''))
                        effect = action.get('effect', action.get('description', ''))
                        action_str = f"{i}. {action_text}"
                        if effect:
                            action_str += f" - {effect}"
                    else:
                        action_str = f"{i}. {action}"
                    
                    ws.merge_cells(f'A{row}:D{row}')
                    ws[f'A{row}'] = action_str
                    ws[f'A{row}'].alignment = wrap_align
                    row += 1
            
            priority_actions = recommendations.get('priority_actions', [])
            if priority_actions:
                row += 1
                ws[f'A{row}'] = "🚨 Срочные действия:"
                ws[f'A{row}'].font = Font(bold=True)
                row += 1
                
                for action in priority_actions:
                    if isinstance(action, dict):
                        action_text = action.get('action', action.get('title', ''))
                    else:
                        action_text = str(action)
                    
                    ws.merge_cells(f'A{row}:D{row}')
                    ws[f'A{row}'] = f"• {action_text}"
                    ws[f'A{row}'].alignment = wrap_align
                    row += 1
        
        # Настройка ширины столбцов
        ws.column_dimensions['A'].width = 25
        ws.column_dimensions['B'].width = 30
        ws.column_dimensions['C'].width = 30
        ws.column_dimensions['D'].width = 30
        
        # Футер
        row += 2
        ws.merge_cells(f'A{row}:D{row}')
        ws[f'A{row}'] = "Сгенерировано DocScan AI - https://docscan-ai.ru"
        ws[f'A{row}'].font = Font(size=9, color="6C757D")
        ws[f'A{row}'].alignment = center_align
        
        # Сохраняем в BytesIO
        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        
        logger.info(f"✅ Excel файл успешно сгенерирован для файла: {filename}")
        return buffer.getvalue()
        
    except Exception as e:
        logger.error(f"❌ Ошибка генерации Excel файла: {e}")
        import traceback
        logger.error(traceback.format_exc())
        raise

